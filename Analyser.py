##########################
# document analyser      #
# todo:                  #
# multiple word lists    #
# multiple words         #
#   
##########################
from docx import Document
import re 

def extract_occurrences(doc_path, target_words):
    occurrences = {word: [] for word in target_words}

    # Load the Word document
    doc = Document(doc_path)

    for paragraph in doc.paragraphs:
            text = paragraph.text
            for word in target_words:
                # Use regular expressions to split the paragraph into sentences
                sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', text)
                for sentence in sentences:
                    if word.lower() in sentence.lower():
                        occurrences[word].append(sentence)

    return occurrences

def save_to_word(result, output_path):
    doc = Document()

    for word, word_occurrences in result.items():
        doc.add_heading(f"Occurrences of '{word}':", level=1)
        for i, occurrence in enumerate(word_occurrences, 1):
            doc.add_paragraph(f"  Occ {i}: {occurrence}")

    doc.save(output_path)

if __name__ == "__main__":

    document_path = "C:/Users/Ben/Desktop/die Weiten/diary copy/copy of diary.docx"
    output_path = "C:/Users/Ben/Desktop/die Weiten/diary copy/resultdoc.docx"

    # Specify the target word
    target_words = ["feel", "think", "important"]

    # Extract occurrences of the target word
    result = extract_occurrences(document_path, target_words)

    save_to_word(result, output_path)
    print(f"Result saved to {output_path}")


