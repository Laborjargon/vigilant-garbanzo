##########################
# document analyser      #
# todo:                  #
# multiple word lists    #
# multiple words         #
#   
##########################
from docx import Document
import re
import pandas as pd

def extract_occurrences(doc_path, target_words):
    occurrences = {word: [] for word in target_words}

    # Load the Word document
    doc = Document(doc_path)

    for paragraph in doc.paragraphs:
            text = paragraph.text
            for word in target_words:
                # Use regular expressions to split the paragraph into sentences
                sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', text)
                for sentence in sentences:
                    if word.lower() in sentence.lower():
                        occurrences[word].append(sentence)

    return occurrences

def count_occurrences(doc_path, target_words):
    counts = {word: 0 for word in target_words}

    # Load the Word document
    doc = Document(doc_path)

    # Iterate through paragraphs and count occurrences of the target words
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for word in target_words:
            # Use regular expressions to split the paragraph into sentences
            sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', text)
            for sentence in sentences:
                if word.lower() in sentence.lower():
                    counts[word] += 1

    return counts

def counts_to_dataframe(counts):
    # Convert counts to a DataFrame
    df = pd.DataFrame(list(counts.items()), columns=['Word', 'Occurrences'])

    return df

def save_to_dataframe(occurrences):
    # Convert occurrences to a DataFrame
    max_length = max(len(sentences) for sentences in occurrences.values())

    # Pad shorter lists with None
    for word, sentences in occurrences.items():
        occurrences[word] += [None] * (max_length - len(sentences))

    # Convert occurrences to a DataFrame
    df = pd.DataFrame(occurrences)

    return df

def save_to_dat(df, output_path):
    # Save the DataFrame to a .dat file
    df.to_csv(output_path, sep='\t', index=False)

def save_to_word(result, output_path):
    doc = Document()

    for word, word_occurrences in result.items():
        doc.add_heading(f"Occurrences of '{word}':", level=1)
        for i, occurrence in enumerate(word_occurrences, 1):
            doc.add_paragraph(f"  Occ {i}: {occurrence}")

    doc.save(output_path)

if __name__ == "__main__":

    document_path = "C:/Users/somme/Desktop/Diary/Mobile Diary.docx"
    output_path = "C:/Users/somme/Desktop/Diary/resultsdoc.docx"
    output_path_dat = "C:/Users/somme/Desktop/Diary/worddata.dat"

    # Specify the target word
    target_words = ["feel", "think", "important"]

    # Extract occurrences of the target word
    result = extract_occurrences(document_path, target_words)

    occurrences = extract_occurrences(document_path, target_words) # lol double

    counts = count_occurrences(document_path, target_words)

    save_to_word(result, output_path)
    
    df = save_to_dataframe(occurrences)
    cdf = counts_to_dataframe(counts)

    # Save the DataFrame to a .dat file
    save_to_dat(df, output_path_dat)
    save_to_dat(cdf, output_path_dat)


    print(f"Result saved to {output_path} and {output_path_dat}")

